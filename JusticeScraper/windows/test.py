import requests
import docx
import unicodedata
import tkinter as tk
import urllib.request
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def get_subjektID(ico):
    url_0 = f"https://or.justice.cz/ias/ui/rejstrik-$firma?ico={ico}&jenPlatne=PLATNE&polozek=1&typHledani=STARTS_WITH"
    html_subjektID = requests.get(url_0).content
    soupID = BeautifulSoup(html_subjektID, 'lxml', from_encoding="utf-8")
    vypis = str(soupID.find(href=True, text='Výpis platných'))
    subjektID = vypis.split('subjektId=')[1].split('&')[0]
    return subjektID

def get_soup(subjektID):
    url = f'https://or.justice.cz/ias/ui/rejstrik-firma.vysledky?subjektId={subjektID}&typ=PLATNY'
    html = requests.get(url).content
    soup = BeautifulSoup(html, 'lxml', from_encoding="utf-8")
    return soup

def get_info(subjektID):
    soup = get_soup(subjektID)
    data = {}
    # first getting the basic information which will be formatted differently
    nazev = soup.find(class_="nounderline").findNext(class_="nounderline").findNext('span').text    #here we take different approach as the object is sometimes Obchodní firma and sometimes different
    data['Název společnosti:'] = nazev

    date = soup.find(text="Datum vzniku a zápisu:").findNext('div').findNext('div').text
    data['Datum vzniku:'] = date

    spis = soup.find(text="Spisová značka: ").findNext('span').text
    data['Spisová značka:'] = spis

    sidlo = soup.find(text="Sídlo: ").findNext('span').findNext('span').text                      #address is in double span
    data['Sídlo:'] = sidlo

    data['IČO:'] = str(ico)

    pravni_forma = soup.find(text="Právní forma: ").findNext('span').text
    data['Právní forma:'] = pravni_forma

    vr_childs = soup.find(text="Právní forma: ").find_all_next(class_='vr-child')  #now the functions looks for all other types of information
    while vr_childs[0].find(class_='nounderline').text == '':    #this part makes sure the code will not get stuck on some companies, which have additional info bellow právní forma, i.e. 45359326 
        vr_childs = vr_childs[1:]
    keys = []
    space = '    '    #this will create indentation later on for subcategories
    for child in vr_childs:
        test = child.find(class_='aunp-udajPanel') #check if there is any information available, sometimes there are empty childs, the code will skip them
        if test is None:
            continue
        all_spans = child.find(class_='aunp-udajPanel').findAll('span')
        parents_vrchild =  child.find_parents(class_='vr-child') 
        spans = []
        for span in all_spans: #getting all text spans that are not child of any other span
            span_child = span.findChildren()
            if span_child:
                pass
            else:
                if span.text.strip(): 
                    spans.append(span)
                else: pass
        first_span = spans[0]
        try:
            if first_span['class'][0] == 'nounderline':
                key = len(parents_vrchild)*space + first_span.text.strip()
                while key in keys:
                    key = key + '+'
                spans = spans[1:]
            else: key = keys[-1]
        except: key = keys[-1]
        i = 0  #set an index which helps keep track of the spans used in the next function
        info = []
        while i < len(spans):  #this loop makes sure, that the information is kept within the same line as shown on justice.cz
            span = spans[i]
            text = span.text
            try:
                span_next = spans[i+1]
            except:
                info.append(unicodedata.normalize("NFKD",text))
                break
            if span.parent == span_next.parent:
                while span.parent == span_next.parent:
                    if str(span.next.next) == '<br/>':   #checks if there are no line breaks inserted
                        break
                    text = text + span_next.text
                    i += 1
                    span = spans[i]
                    try:
                        span_next = spans[i+1]
                    except:
                        break
                info.append(unicodedata.normalize("NFKD",text))
                i += 1
            else:
                info.append(unicodedata.normalize("NFKD",text))
                i += 1
        if key in keys:
            data[key].extend(info)
        else:
            data[key] = info
        keys.append(key)
    return data

def get_vypis_doc(subjektID):
    data = get_info(subjektID)
    keys = list(data.keys())

    #define document and styles settings
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.first_line_indent = Inches(-1.4)
    style.paragraph_format.left_indent = Inches(1.4)
    style.paragraph_format.space_after = Inches(0)
    style.paragraph_format.space_before = Inches(0)

    styles = document.styles
    style1 = styles.add_style('Light', WD_STYLE_TYPE.CHARACTER)
    font1 = style1.font
    font1.name = 'Calibri Light'
    font1.size = Pt(11)

    #get the basic info, which is displayed on the same line:
    #název společnosti
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[0]).bold=True
    p.add_run('\t'+data[keys[0]]).bold=True
    #datum vzniku
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[1]).bold=True
    p.add_run('\t'+data[keys[1]], style=document.styles['Light'])
    #spisová značka
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[2]).bold=True
    p.add_run('\t'+data[keys[2]], style=document.styles['Light'])
    #sídlo
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[3]).bold=True
    p.add_run('\t'+data[keys[3]], style=document.styles['Light'])
    #ičo
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[4]).bold=True
    p.add_run('\t'+data[keys[4]], style=document.styles['Light'])
    #právní forma
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[5]).bold=True
    p.add_run('\t'+data[keys[5]], style=document.styles['Light'])
    #get the rest of the information
    for key in keys[6:]: 
        p = document.add_paragraph(style=document.styles['Normal'])
        p.add_run(key.replace('+','')).bold=True
        values = data[key]
        for value in values:
            p.add_run('\n'+value, style=document.styles['Light'])

    document.save(f'výpis_{data[keys[0]]}_.docx')
    return data[keys[0]]

def get_pdf_file(subjektID, nazev):
    pdf_url = f'https://or.justice.cz/ias/ui/print-pdf?subjektId={subjektID}&typVypisu=PLATNY&full=false'
    pdf_filename = f'pdf_výpis_{nazev}_.pdf'
    urllib.request.urlretrieve(pdf_url, pdf_filename)

def get_vypis(ico):
    subjektID = get_subjektID(ico)
    try:
        nazev = get_vypis_doc(subjektID)
        get_pdf_file(subjektID, nazev)
        print('done')
    except: 
        print('error')

ico = 27201121
get_vypis(ico)