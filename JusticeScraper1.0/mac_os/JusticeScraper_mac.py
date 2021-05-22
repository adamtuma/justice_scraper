#JUSTICE SCRAPER 1.0 macOS
import requests
import docx
import unicodedata
import tkinter as tk
import os
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def get_url(ico):
    url_0 = f'https://or.justice.cz/ias/ui/rejstrik-$firma?ico={ico}&jenPlatne=PLATNE&polozek=1&typHledani=STARTS_WITH'
    html_subjektID = requests.get(url_0).content
    soupID = BeautifulSoup(html_subjektID, 'lxml', from_encoding="utf-8")
    vypis = str(soupID.find(href=True, text='Výpis platných'))
    subjektID = vypis.split('subjektId=')[1].split('&')[0]
    url = f'https://or.justice.cz/ias/ui/rejstrik-firma.vysledky?subjektId={subjektID}&typ=PLATNY'
    return url

def get_soup(ico):
    url = get_url(ico)
    html = requests.get(url).content
    soup = BeautifulSoup(html, 'lxml', from_encoding="utf-8")
    return soup

def get_info(ico):
    soup = get_soup(ico)
    data = {}
    # first getting the basic information which will be formatted differently
    nazev = soup.find(class_="nounderline").findNext(class_="nounderline").findNext('span').text    #here we take different approach as the object is sometimes Obchodní firma and sometimes different
    data['Název společnosti:'] = nazev

    date = soup.find(text="Datum vzniku a zápisu:").findNext('div').findNext('div').text
    data['Datum vzniku:'] = date

    spis = soup.find(text="Spisová značka: ").findNext('span').text
    data['Spisová značka:'] = spis

    sidlo = soup.find(text="Sídlo: ").findNext('span').findNext('span').text                      #address is in double span
    data['Sídlo:'] = sidlo

    data['IČO:'] = str(ico)

    pravni_forma = soup.find(text="Právní forma: ").findNext('span').text
    data['Právní forma:'] = pravni_forma

    vr_childs = soup.find(text="Právní forma: ").find_all_next(class_='vr-child')  #now the functions looks for all other types of information
    while vr_childs[0].find(class_='nounderline').text == '':    #this part makes sure the code will not get stuck on some companies, which have additional info bellow právní forma, i.e. 45359326 
        vr_childs = vr_childs[1:]
    keys = []
    space = '    '    #this will create indentation later on for subcategories
    for child in vr_childs:
        test = child.find(class_='aunp-udajPanel') #check if there is any information available, sometimes there are empty childs, the code will skip them
        if test is None:
            continue
        all_spans = child.find(class_='aunp-udajPanel').findAll('span')
        parents_vrchild =  child.find_parents(class_='vr-child') 
        spans = []
        for span in all_spans: #getting all text spans that are not child of any other span
            span_child = span.findChildren()
            if span_child:
                pass
            else:
                if span.text.strip(): 
                    spans.append(span)
                else: pass
        first_span = spans[0]
        try:
            if first_span['class'][0] == 'nounderline':
                key = len(parents_vrchild)*space + first_span.text.strip()
                while key in keys:
                    key = key + '+'
                spans = spans[1:]
            else: key = keys[-1]
        except: key = keys[-1]
        i = 0  #set an index which helps keep track of the spans used in the next function
        info = []
        while i < len(spans):  #this loop makes sure, that the information is kept within the same line as shown on justice.cz
            span = spans[i]
            text = span.text
            try:
                span_next = spans[i+1]
            except:
                info.append(unicodedata.normalize("NFKD",text))
                break
            if span.parent == span_next.parent:
                while span.parent == span_next.parent:
                    if str(span.next.next) == '<br/>':   #checks if there are no line breaks inserted
                        break
                    text = text + span_next.text
                    i += 1
                    span = spans[i]
                    try:
                        span_next = spans[i+1]
                    except:
                        break
                info.append(unicodedata.normalize("NFKD",text))
                i += 1
            else:
                info.append(unicodedata.normalize("NFKD",text))
                i += 1
        if key in keys:
            data[key].extend(info)
        else:
            data[key] = info
        keys.append(key)
    return data

def get_vypis_doc(ico):
    data = get_info(ico)
    keys = list(data.keys())

    #define document and styles settings
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.first_line_indent = Inches(-1.4)
    style.paragraph_format.left_indent = Inches(1.4)
    style.paragraph_format.space_after = Inches(0)
    style.paragraph_format.space_before = Inches(0)

    styles = document.styles
    style1 = styles.add_style('Light', WD_STYLE_TYPE.CHARACTER)
    font1 = style1.font
    font1.name = 'Calibri Light'
    font1.size = Pt(11)

    #get the basic info, which is displayed on the same line:
    #název společnosti
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[0]).bold=True
    p.add_run('\t'+data[keys[0]]).bold=True
    #datum vzniku
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[1]).bold=True
    p.add_run('\t'+data[keys[1]], style=document.styles['Light'])
    #spisová značka
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[2]).bold=True
    p.add_run('\t'+data[keys[2]], style=document.styles['Light'])
    #sídlo
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[3]).bold=True
    p.add_run('\t'+data[keys[3]], style=document.styles['Light'])
    #ičo
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[4]).bold=True
    p.add_run('\t'+data[keys[4]], style=document.styles['Light'])
    #právní forma
    p = document.add_paragraph(style=document.styles['Normal'])
    p.add_run(keys[5]).bold=True
    p.add_run('\t'+data[keys[5]], style=document.styles['Light'])
    #get the rest of the information
    if var1.get() == 0:
        if 'Ostatní skutečnosti:' in keys:  #not interested in 'ostatní skutečnosti'
            position = keys.index('Ostatní skutečnosti:')
            keys = keys[:position]
    for key in keys[6:]: 
        p = document.add_paragraph(style=document.styles['Normal'])
        p.add_run(key.replace('+','')).bold=True
        values = data[key]
        for value in values:
            p.add_run('\n'+value, style=document.styles['Light'])

    location = os.path.abspath(
        os.path.join(sys.executable + f'/výpis_{data[keys[0]]}_.docx', '..', '..', '..', '..','..', f'výpis_{data[keys[0]]}_.docx'))
    #location = '/Users/adamtuma/Documents/mac_os/'+f'/výpis_{data[keys[0]]}_.docx'
    document.save(location)

root= tk.Tk()

canvas1 = tk.Canvas(root, width = 400, height = 300)
canvas1.pack()

label1 = tk.Label(root, text='Justice Scraper 1.0')
label1.config(font=('Calibri', 14))
canvas1.create_window(200, 25, window=label1)

label2 = tk.Label(root, text='IČO společnosti:')
label2.config(font=('Calibri', 10))
canvas1.create_window(200, 120, window=label2)

label3 = tk.Label(root, text=('Adam Tůma 2021'+'\u00A9'))
label3.config(font=('Calibri', 9))
canvas1.create_window(55, 290, window=label3)

entry1 = tk.Entry (root) 
canvas1.create_window(200, 140, window=entry1)

var1 = tk.IntVar()
var1.set(1)
checker1 = tk.Checkbutton(root, text='Ostatní skutečnosti',variable=var1, onvalue=1, offvalue=0)
canvas1.create_window(200, 165, window=checker1)

label4 = tk.Label(root)
canvas1.create_window(200, 230, window=label4)

def get_vypis():
    ico = str(entry1.get())
    label4.config(text='')
    try:
        get_vypis_doc(ico)
        label4.config(text='Hotovo!')
    except: 
        label4.config(text='Něco se pokazilo, zkontroluj IČO.')

button1 = tk.Button(text='Připravit výpis', command=get_vypis)
button1.config(font=('Calibri', 10))
canvas1.create_window(200, 195, window=button1)

root.mainloop()